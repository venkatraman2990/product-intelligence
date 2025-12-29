"""Document loader for PDF and Word files.

Provides a unified interface for extracting text from various document formats.
"""

import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Container for loaded document data."""

    path: Path
    text: str
    page_count: int
    document_type: str
    metadata: dict


class DocumentLoader:
    """Load and extract text from PDF and Word documents."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

    def __init__(self):
        """Initialize the document loader."""
        self._check_dependencies()

    def _check_dependencies(self):
        """Check if required libraries are available."""
        try:
            import fitz  # PyMuPDF

            self._has_pymupdf = True
        except ImportError:
            self._has_pymupdf = False
            logger.warning("PyMuPDF not installed. PDF support disabled.")

        try:
            import docx

            self._has_docx = True
        except ImportError:
            self._has_docx = False
            logger.warning("python-docx not installed. Word support disabled.")

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load a document and extract its text content.

        Args:
            file_path: Path to the document file

        Returns:
            LoadedDocument with extracted text and metadata

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        if extension == ".pdf":
            return self._load_pdf(path)
        elif extension in {".docx", ".doc"}:
            return self._load_word(path)

    def _load_pdf(self, path: Path) -> LoadedDocument:
        """Extract text from a PDF file using PyMuPDF."""
        if not self._has_pymupdf:
            raise RuntimeError(
                "PyMuPDF is required for PDF support. "
                "Install with: pip install PyMuPDF"
            )

        import fitz  # PyMuPDF

        logger.info(f"Loading PDF: {path}")

        doc = fitz.open(path)
        total_pages = doc.page_count  # Get actual PDF page count
        text_parts = []
        metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "creation_date": doc.metadata.get("creationDate", ""),
        }

        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

        doc.close()

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from {total_pages} pages ({len(text_parts)} with text)")

        return LoadedDocument(
            path=path,
            text=full_text,
            page_count=total_pages,  # Use actual page count, not just pages with text
            document_type="pdf",
            metadata=metadata,
        )

    def _load_word(self, path: Path) -> LoadedDocument:
        """Extract text from a Word document using python-docx."""
        if not self._has_docx:
            raise RuntimeError(
                "python-docx is required for Word support. "
                "Install with: pip install python-docx"
            )

        from docx import Document

        logger.info(f"Loading Word document: {path}")

        doc = Document(path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n[Table]\n" + "\n".join(table_text))

        # Get metadata from core properties
        metadata = {}
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception as e:
            logger.warning(f"Could not extract Word metadata: {e}")

        full_text = "\n\n".join(text_parts)

        # Estimate page count based on content length (~3000 chars per page typical)
        estimated_pages = max(1, len(full_text) // 3000 + (1 if len(full_text) % 3000 > 500 else 0))

        logger.info(f"Extracted {len(full_text)} characters from Word document (~{estimated_pages} pages)")

        return LoadedDocument(
            path=path,
            text=full_text,
            page_count=estimated_pages,  # Estimated based on content length
            document_type="docx",
            metadata=metadata,
        )

    def load_batch(self, file_paths: list[str | Path]) -> list[LoadedDocument]:
        """Load multiple documents.

        Args:
            file_paths: List of paths to documents

        Returns:
            List of LoadedDocument objects (skipping failed loads)
        """
        results = []
        for path in file_paths:
            try:
                doc = self.load(path)
                results.append(doc)
            except Exception as e:
                logger.error(f"Failed to load {path}: {e}")
        return results

    @classmethod
    def is_supported(cls, file_path: str | Path) -> bool:
        """Check if a file type is supported.

        Args:
            file_path: Path to check

        Returns:
            True if the file extension is supported
        """
        return Path(file_path).suffix.lower() in cls.SUPPORTED_EXTENSIONS
